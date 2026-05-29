import os
from pathlib import Path
from datetime import datetime

DEFAULT_SAVE_PATH = str(Path.home() / "Documents")


def create_word(filename: str, title: str, content: str, save_path: str = None) -> dict:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"error": "Module python-docx manquant. Lance : pip install python-docx"}

    save_dir = Path(save_path or DEFAULT_SAVE_PATH).expanduser()
    save_dir.mkdir(parents=True, exist_ok=True)
    filepath = save_dir / f"{filename}.docx"

    try:
        doc = Document()

        # Titre
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(datetime.now().strftime("%d/%m/%Y"))
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()

        # Contenu — on gère les sauts de ligne
        for paragraph in content.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
            else:
                doc.add_paragraph()

        doc.save(str(filepath))
        return {
            "success":  True,
            "fichier":  filepath.name,
            "chemin":   str(filepath),
            "taille":   f"{filepath.stat().st_size / 1024:.1f} Ko",
        }
    except Exception as e:
        return {"error": str(e)}


def create_txt(filename: str, content: str, save_path: str = None) -> dict:
    save_dir = Path(save_path or DEFAULT_SAVE_PATH).expanduser()
    save_dir.mkdir(parents=True, exist_ok=True)
    filepath = save_dir / f"{filename}.txt"

    try:
        filepath.write_text(content, encoding="utf-8")
        return {
            "success": True,
            "fichier": filepath.name,
            "chemin":  str(filepath),
        }
    except Exception as e:
        return {"error": str(e)}


def create_pdf(filename: str, title: str, content: str, save_path: str = None) -> dict:
    try:
        from fpdf import FPDF
    except ImportError:
        return {"error": "Module fpdf2 manquant. Lance : pip install fpdf2"}

    save_dir = Path(save_path or DEFAULT_SAVE_PATH).expanduser()
    save_dir.mkdir(parents=True, exist_ok=True)
    filepath = save_dir / f"{filename}.pdf"

    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)

        # Titre
        pdf.set_font("Helvetica", style="B", size=16)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT", align="C")

        # Date
        pdf.set_font("Helvetica", size=9)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(0, 6, datetime.now().strftime("%d/%m/%Y"), new_x="LMARGIN", new_y="NEXT", align="R")
        pdf.ln(4)

        # Contenu
        pdf.set_font("Helvetica", size=11)
        pdf.set_text_color(40, 40, 40)

        for line in content.split("\n"):
            if line.strip():
                pdf.multi_cell(0, 6, line)
            else:
                pdf.ln(3)

        pdf.output(str(filepath))
        return {
            "success": True,
            "fichier": filepath.name,
            "chemin":  str(filepath),
            "taille":  f"{filepath.stat().st_size / 1024:.1f} Ko",
        }
    except Exception as e:
        return {"error": str(e)}


def read_document(path: str) -> dict:
    p = Path(path).expanduser()
    if not p.exists():
        return {"error": f"Fichier introuvable : {path}"}

    ext = p.suffix.lower()

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(str(p))
            text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            if len(text) > 8000:
                text = text[:8000] + "\n\n[... tronqué ...]"
            return {"success": True, "nom": p.name, "type": "Word", "contenu": text}
        except ImportError:
            return {"error": "Module python-docx manquant. Lance : pip install python-docx"}
        except Exception as e:
            return {"error": str(e)}

    elif ext == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(str(p))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if len(text) > 8000:
                text = text[:8000] + "\n\n[... tronqué ...]"
            return {"success": True, "nom": p.name, "type": "PDF", "pages": len(reader.pages), "contenu": text}
        except ImportError:
            return {"error": "Module pypdf manquant. Lance : pip install pypdf"}
        except Exception as e:
            return {"error": str(e)}

    elif ext == ".txt":
        try:
            text = p.read_text(encoding="utf-8", errors="replace")
            if len(text) > 8000:
                text = text[:8000] + "\n\n[... tronqué ...]"
            return {"success": True, "nom": p.name, "type": "TXT", "contenu": text}
        except Exception as e:
            return {"error": str(e)}

    else:
        return {"error": f"Format non supporté : {ext}. Supporte : .docx, .pdf, .txt"}
