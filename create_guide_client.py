"""
Génère le guide d'installation client — document Word professionnel
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

GOLD  = RGBColor(0xC8, 0xA9, 0x6E)
DARK  = RGBColor(0x1A, 0x1A, 0x17)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY  = RGBColor(0x77, 0x77, 0x77)
RED   = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x27, 0xAE, 0x60)


def cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def titre(doc, text, level=1):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size      = Pt(18)
        run.font.color.rgb = GOLD
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        run.font.size      = Pt(13)
        run.font.color.rgb = DARK
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    return p


def para(doc, text, bold=False, italic=False, color=None, size=11, align=None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    return p


def etape(doc, numero, titre_etape, description):
    """Bloc étape numérotée avec fond doré clair"""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Numéro
    num_cell = table.rows[0].cells[0]
    num_cell.width = Inches(0.5)
    cell_bg(num_cell, 'C8A96E')
    p = num_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(numero))
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = WHITE

    # Contenu
    txt_cell = table.rows[0].cells[1]
    cell_bg(txt_cell, 'F9F6F0')
    p2 = txt_cell.paragraphs[0]
    r1 = p2.add_run(titre_etape + "\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK
    r2 = p2.add_run(description)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GREY

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def encadre(doc, text, couleur='F0EBE0', bordure='C8A96E'):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    c = table.rows[0].cells[0]
    cell_bg(c, couleur)
    p = c.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def bullet(doc, items, color=None):
    for item in items:
        p   = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        if color:
            run.font.color.rgb = color
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3)


def separateur(doc):
    p = doc.add_paragraph("─" * 72)
    p.paragraphs if False else None
    run = p.runs[0]
    run.font.color.rgb = GOLD
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    # ══════════════════════════════════════════════════════
    # EN-TÊTE
    # ══════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("THE AMAH")
    r.bold = True
    r.font.size = Pt(32)
    r.font.color.rgb = GOLD

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Agent IA Local — Guide d'installation")
    r2.font.size = Pt(14)
    r2.font.color.rgb = GREY

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"contact.amah.officiel@gmail.com  •  {datetime.now().strftime('%d/%m/%Y')}")
    r3.font.size = Pt(10)
    r3.font.color.rgb = GREY
    r3.italic = True

    separateur(doc)

    # ══════════════════════════════════════════════════════
    # BIENVENUE
    # ══════════════════════════════════════════════════════
    titre(doc, "Bienvenue avec Amah", 1)

    para(doc,
        "Amah est votre assistante IA personnelle installée directement sur votre PC. "
        "Elle peut gérer vos fichiers, créer des documents, lire vos emails, "
        "naviguer sur internet et bien plus encore — le tout en français, "
        "simplement en lui parlant.")

    para(doc,
        "Ce guide vous accompagne pas à pas pour installer et configurer Amah "
        "sur votre ordinateur. Suivez les étapes dans l'ordre et tout se passera bien.",
        italic=True, color=GREY)

    separateur(doc)

    # ══════════════════════════════════════════════════════
    # CE QUE VOUS AVEZ REÇU
    # ══════════════════════════════════════════════════════
    titre(doc, "Ce que vous avez reçu", 1)

    para(doc, "Votre dossier d'installation contient 4 fichiers :")

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = table.rows[0].cells
    headers[0].text = "Fichier"
    headers[1].text = "À quoi ça sert"
    for c in headers:
        cell_bg(c, '1A1A17')
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = WHITE

    contenu = [
        ("Amah Agent.exe",           "Le programme principal — c'est lui que vous lancez"),
        (".env",                      "Votre fichier de configuration (ne pas supprimer)"),
        ("installer_navigateur.bat",  "Installe Chrome pour la navigation web (une seule fois)"),
        ("GUIDE_INSTALLATION.md",     "Ce guide en version texte"),
    ]
    for i, (fichier, desc) in enumerate(contenu):
        row = table.rows[i+1].cells
        r1 = row[0].paragraphs[0].add_run(fichier)
        r1.bold = True
        r1.font.color.rgb = GOLD
        row[1].text = desc

    doc.add_paragraph()
    separateur(doc)

    # ══════════════════════════════════════════════════════
    # ÉTAPE 1 — PLACER LES FICHIERS
    # ══════════════════════════════════════════════════════
    titre(doc, "Installation — Étape par étape", 1)

    etape(doc, 1, "Placez les fichiers au bon endroit",
          "Créez un dossier nommé « Amah Agent » sur votre bureau ou dans vos Documents.\n"
          "Copiez les 4 fichiers reçus dans ce dossier. Ne les dispersez pas.")

    # ══════════════════════════════════════════════════════
    # ÉTAPE 2 — OBTENIR LA CLÉ GROQ
    # ══════════════════════════════════════════════════════
    etape(doc, 2, "Créez votre clé API Groq (gratuit, 2 minutes)",
          "Groq est le service IA qui donne son intelligence à Amah. Il est entièrement gratuit.")

    titre(doc, "Comment créer votre clé Groq :", 2)
    bullet(doc, [
        "Ouvrez votre navigateur et allez sur : console.groq.com",
        "Cliquez sur « Sign Up » et créez un compte gratuit (ou connectez-vous)",
        "Dans le menu à gauche, cliquez sur « API Keys »",
        "Cliquez sur « Create API Key »",
        "Donnez-lui un nom (exemple : Amah) et cliquez sur « Submit »",
        "Copiez la clé affichée — elle commence par « gsk_ »",
        "⚠️ Conservez-la précieusement, elle ne sera plus affichée après",
    ])

    encadre(doc,
        "💡 Limites du compte gratuit : 30 requêtes par minute, 14 400 par jour.\n"
        "C'est largement suffisant pour un usage quotidien normal.")

    # ══════════════════════════════════════════════════════
    # ÉTAPE 3 — PREMIER LANCEMENT
    # ══════════════════════════════════════════════════════
    etape(doc, 3, "Premier lancement de Amah Agent",
          "Double-cliquez sur « Amah Agent.exe » dans votre dossier d'installation.")

    para(doc,
        "Si Windows affiche un avertissement « Application inconnue » → "
        "cliquez sur « Informations complémentaires » puis « Exécuter quand même ».",
        color=GREY, italic=True)

    para(doc,
        "Une fenêtre de configuration s'ouvre automatiquement :",
        bold=True)

    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Table Grid'

    h2 = table2.rows[0].cells
    h2[0].text = "Champ"
    h2[1].text = "Que mettre"
    for c in h2:
        cell_bg(c, '1A1A17')
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = WHITE

    lignes = [
        ("Clé API Groq (obligatoire)",        "Collez la clé copiée à l'étape 2 (commence par gsk_...)"),
        ("Adresse Gmail (optionnel)",          "Laissez tel quel — déjà configurée pour vous"),
        ("Mot de passe Gmail (optionnel)",     "Laissez vide si non communiqué par votre prestataire"),
    ]
    for i, (champ, val) in enumerate(lignes):
        row = table2.rows[i+1].cells
        r = row[0].paragraphs[0].add_run(champ)
        r.bold = True
        row[1].text = val

    doc.add_paragraph()
    para(doc, "→ Cliquez sur « Démarrer Amah » — votre configuration est sauvegardée automatiquement.",
         bold=True, color=GREEN)

    encadre(doc,
        "✅ À partir de maintenant, Amah se souvient de votre configuration.\n"
        "Les prochains lancements démarrent directement sans passer par cet écran.")

    # ══════════════════════════════════════════════════════
    # ÉTAPE 4 — UTILISATION
    # ══════════════════════════════════════════════════════
    etape(doc, 4, "Utilisation quotidienne",
          "Double-cliquez sur « Amah Agent.exe » pour lancer Amah. Parlez-lui en français.")

    titre(doc, "Exemples de ce que vous pouvez lui demander :", 2)

    exemples = doc.add_table(rows=6, cols=2)
    exemples.style = 'Table Grid'

    h3 = exemples.rows[0].cells
    h3[0].text = "Ce que vous dites"
    h3[1].text = "Ce qu'Amah fait"
    for c in h3:
        cell_bg(c, 'C8A96E')
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = WHITE

    ex = [
        ("Organise mon bureau automatiquement",     "Classe tous les fichiers par type (Images, Documents...)"),
        ("Lis mes 5 derniers emails",               "Affiche les expéditeurs, sujets et extraits"),
        ("Crée un rapport Word sur notre réunion",  "Génère un document .docx complet"),
        ("Cherche le prix de l'iPhone 16",          "Recherche sur internet et répond"),
        ("Mémorise que je préfère les PDF",         "Retient cette préférence pour toujours"),
    ]
    for i, (q, r_) in enumerate(ex):
        row = exemples.rows[i+1].cells
        rq = row[0].paragraphs[0].add_run(q)
        rq.italic = True
        row[1].text = r_

    doc.add_paragraph()
    separateur(doc)

    # ══════════════════════════════════════════════════════
    # ÉTAPE 5 — NAVIGATEUR WEB (OPTIONNEL)
    # ══════════════════════════════════════════════════════
    titre(doc, "Étape optionnelle — Navigation web automatique", 1)

    para(doc,
        "Si vous souhaitez qu'Amah puisse naviguer sur internet avec Chrome visible "
        "(ouvrir des pages, remplir des formulaires, faire des captures d'écran), "
        "vous devez installer le navigateur une seule fois.")

    etape(doc, "●", "Installer le navigateur Chrome pour Amah",
          "Double-cliquez sur « installer_navigateur.bat » dans votre dossier.\n"
          "Une fenêtre noire s'ouvre et télécharge Chrome automatiquement (~170 Mo).\n"
          "Attendez que ce soit terminé, puis fermez la fenêtre.")

    encadre(doc,
        "⚠️ Cette étape nécessite une connexion internet.\n"
        "Elle ne s'effectue qu'une seule fois — pas besoin de la refaire.")

    separateur(doc)

    # ══════════════════════════════════════════════════════
    # RACCOURCI BUREAU
    # ══════════════════════════════════════════════════════
    titre(doc, "Créer un raccourci sur le bureau (recommandé)", 1)

    para(doc,
        "Pour lancer Amah directement depuis votre bureau sans chercher le fichier :")

    bullet(doc, [
        "Faites un clic droit sur « Amah Agent.exe »",
        "Choisissez « Envoyer vers » → « Bureau (créer un raccourci) »",
        "Une icône Amah apparaît sur votre bureau — double-cliquez pour lancer",
    ])

    separateur(doc)

    # ══════════════════════════════════════════════════════
    # PROBLÈMES FRÉQUENTS
    # ══════════════════════════════════════════════════════
    titre(doc, "Problèmes fréquents et solutions", 1)

    prob = doc.add_table(rows=6, cols=2)
    prob.style = 'Table Grid'

    hp = prob.rows[0].cells
    hp[0].text = "Problème"
    hp[1].text = "Solution"
    for c in hp:
        cell_bg(c, '1A1A17')
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = WHITE

    pbs = [
        ("Windows bloque le .exe au lancement",
         "Clic droit → « Propriétés » → cochez « Débloquer » → OK. Puis relancez."),
        ("« Clé Groq invalide » au démarrage",
         "Vérifiez que vous avez bien copié TOUTE la clé (commence par gsk_...). Relancez."),
        ("L'écran de config réapparaît à chaque lancement",
         "Le fichier .env a été supprimé. Recopiez-le depuis votre dossier d'installation."),
        ("Amah ne trouve pas mes emails",
         "Vérifiez que le fichier .env est bien dans le même dossier que le .exe."),
        ("Le navigateur ne s'ouvre pas",
         "Lancez d'abord « installer_navigateur.bat » une seule fois."),
    ]
    for i, (p_, s) in enumerate(pbs):
        row = prob.rows[i+1].cells
        rp = row[0].paragraphs[0].add_run(p_)
        rp.bold = True
        rp.font.color.rgb = RED
        row[1].text = s

    doc.add_paragraph()
    separateur(doc)

    # ══════════════════════════════════════════════════════
    # CONTACT
    # ══════════════════════════════════════════════════════
    titre(doc, "Besoin d'aide ?", 1)

    para(doc,
        "Pour toute question d'installation, de configuration ou d'utilisation, "
        "contactez-nous par email. Nous répondons sous 24h.")

    p_mail = doc.add_paragraph()
    p_mail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_mail = p_mail.add_run("contact.amah.officiel@gmail.com")
    r_mail.bold = True
    r_mail.font.size = Pt(14)
    r_mail.font.color.rgb = GOLD

    doc.add_paragraph()

    # Pied de page
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run("— THE AMAH — Agent IA Local — Guide Client v1.0 —")
    rf.italic = True
    rf.font.size = Pt(9)
    rf.font.color.rgb = GREY

    # Sauvegarde
    path = r"C:\Users\Smarte technologui\Desktop\Projets\amah-agent\dist\Guide_Installation_Client.docx"
    doc.save(path)

    # Aussi dans le projet
    path2 = r"C:\Users\Smarte technologui\Desktop\Projets\amah-agent\Guide_Installation_Client.docx"
    doc.save(path2)

    print(f"Document cree : {path}")
    os.startfile(path)


if __name__ == "__main__":
    main()
