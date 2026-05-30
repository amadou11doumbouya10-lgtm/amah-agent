"""Génère le document explicatif sur le système de licence Amah Agent."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

GOLD  = RGBColor(0xC8, 0xA9, 0x6E)
DARK  = RGBColor(0x1A, 0x1A, 0x17)
GREY  = RGBColor(0x55, 0x55, 0x55)
RED   = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x27, 0xAE, 0x60)


def h(doc, text, level=1, color=None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(18)
        run.font.color.rgb = color or GOLD
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.color.rgb = color or DARK
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
    return p


def para(doc, text, italic=False, color=None, size=11):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.italic      = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(doc, items, indent=True):
    for item in items:
        p   = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        if indent:
            p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(2)


def box(doc, title, content, color_hex="F5F0E8"):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    # Fond coloré
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

    p1  = cell.paragraphs[0]
    r1  = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK

    p2  = cell.add_paragraph()
    r2  = p2.add_run(content)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GREY
    doc.add_paragraph()


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── TITRE ────────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("THE AMAH — SYSTÈME DE LICENCE")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = GOLD

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("Guide complet — Vendeur et Client")
    rs.font.size = Pt(12)
    rs.font.color.rgb = GREY
    rs.italic = True

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = d.add_run(f"contact.amah.officiel@gmail.com  |  {datetime.now().strftime('%d/%m/%Y')}")
    rd.font.size = Pt(10)
    rd.font.color.rgb = GREY

    doc.add_paragraph()

    # ── 1. QU'EST-CE QU'UNE LICENCE ? ───────────────────────────────────────
    h(doc, "1. Qu'est-ce qu'une licence Amah ?")
    para(doc,
        "Une licence Amah est un code d'activation unique qui permet à un client "
        "d'utiliser Amah Agent sur son PC Windows. Sans licence valide, le logiciel "
        "affiche un écran d'activation et ne démarre pas.")
    para(doc,
        "La licence est liée au PC du client — elle ne peut pas être copiée ou "
        "partagée avec quelqu'un d'autre. Si le client change de PC, il a besoin "
        "d'une nouvelle licence.")

    box(doc,
        "Exemple de clé de licence",
        "B6AD7-18D9B-9210F-DA7FF\n\n"
        "Ce code de 20 caractères est généré spécifiquement pour un PC précis.\n"
        "Il ne fonctionnera sur aucun autre PC.",
        "EFF5E0")

    # ── 2. COMMENT ÇA MARCHE ? ──────────────────────────────────────────────
    h(doc, "2. Comment ça marche (explication simple)")
    para(doc,
        "Chaque PC Windows possède un identifiant unique appelé Machine UUID. "
        "C'est comme le numéro de série de l'ordinateur. "
        "La clé de licence est calculée mathématiquement à partir de cet identifiant.")

    h(doc, "Schéma du processus", level=2)
    steps = [
        "Le client achète Amah Agent et t'envoie son Machine UUID",
        "Tu génères sa clé personnalisée en 30 secondes",
        "Tu lui envoies la clé par email",
        "Il l'entre dans Amah → le logiciel s'active définitivement",
        "La prochaine fois qu'il lance Amah, pas besoin de re-saisir la clé",
    ]
    for i, step in enumerate(steps, 1):
        p   = doc.add_paragraph()
        run = p.add_run(f"  {i}.  {step}")
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # ── 3. CÔTÉ VENDEUR ─────────────────────────────────────────────────────
    h(doc, "3. Ce que TU fais (côté vendeur)")

    h(doc, "Étape 1 — Recevoir le Machine UUID du client", level=2)
    para(doc, "Le client doit ouvrir PowerShell sur son PC et taper la commande :")
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.4)
    cr = code.add_run("(Get-WmiObject Win32_ComputerSystemProduct).UUID")
    cr.font.name  = "Courier New"
    cr.font.size  = Pt(10)
    cr.font.color.rgb = DARK
    para(doc,
        "Il copie le résultat (ex: 4C4C4544-0051-5310-8052-CAC04F525132) et te l'envoie "
        "par email ou WhatsApp.", italic=True, color=GREY)

    h(doc, "Étape 2 — Générer la clé de licence", level=2)
    para(doc, "Une fois que tu as le Machine UUID, tu as deux options :")
    bullet(doc, [
        "Double-clic sur developpeur/generate_license.bat → entre l'UUID → la clé s'affiche",
        "Ou en ligne de commande : py -3.13 tools/license.py <UUID_DU_CLIENT>",
    ])

    h(doc, "Étape 3 — Envoyer la clé au client", level=2)
    para(doc,
        "Tu envoies la clé générée par email. Le client l'entre dans l'écran "
        "d'activation d'Amah. C'est tout.")

    box(doc,
        "Note importante",
        "Garde un tableau de toutes les licences générées :\n"
        "Nom du client | Machine UUID | Clé générée | Date | Version\n\n"
        "C'est indispensable si le client a besoin d'aide ou veut changer de PC.",
        "FFF3CD")

    # ── 4. CÔTÉ CLIENT ──────────────────────────────────────────────────────
    h(doc, "4. Ce que fait le CLIENT")

    h(doc, "Première fois — Activation", level=2)
    bullet(doc, [
        "Il reçoit le fichier Amah Agent.exe et le place dans un dossier",
        "Il double-clique sur Amah Agent.exe",
        "Un écran d'activation s'ouvre automatiquement",
        "Il voit son Machine UUID affiché — il peut le copier directement depuis cet écran",
        "Il entre sa clé de licence dans le champ prévu",
        "Il clique sur « Activer Amah » → le logiciel démarre",
    ])

    h(doc, "Les fois suivantes", level=2)
    para(doc,
        "Le client double-clique simplement sur Amah Agent.exe. "
        "La licence est mémorisée dans le fichier .env — pas besoin de la re-saisir.")

    # ── 5. CAS PARTICULIERS ─────────────────────────────────────────────────
    h(doc, "5. Cas particuliers et FAQ")

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Situation"
    hdr_cells[1].text = "Solution"
    for cell in hdr_cells:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A1A17')
        tcPr.append(shd)

    rows = [
        ("Le client perd sa clé",
         "Regenere-la avec son UUID (garde ton tableau de licences)"),
        ("Le client change de PC",
         "Il te donne le nouveau UUID, tu génères une nouvelle clé\n(compte comme une nouvelle licence ou offre-la selon ta politique)"),
        ("Le client réinstalle Windows",
         "Son UUID reste le même — sa clé fonctionne toujours"),
        ("Le client veut 2 PC",
         "2 UUID différents = 2 licences différentes = 2 achats"),
        ("Tu veux bloquer un client",
         "Publie une nouvelle version avec une nouvelle clé secrète\n(toutes les anciennes clés deviennent invalides)"),
        ("La clé « invalide » alors qu'elle est correcte",
         "Vérifier que le client a bien copié-collé sans espace,\nou qu'il utilise le bon .exe (même version que la clé)"),
    ]
    alt = False
    for q, a in rows:
        row = table.add_row().cells
        row[0].text = q
        row[1].text = a
        fill = "F5F0E8" if alt else "FFFFFF"
        for cell in row:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
        alt = not alt

    doc.add_paragraph()

    # ── 6. SÉCURITÉ ─────────────────────────────────────────────────────────
    h(doc, "6. Sécurité du système")
    bullet(doc, [
        "La clé est liée au hardware — impossible de la transférer sans ton accord",
        "La clé secrète de génération n'est pas dans le logiciel — elle est dans ton .env",
        "Même si quelqu'un décompile le .exe, il ne peut pas générer de clés",
        "Le système fonctionne sans internet — pas de serveur à maintenir",
        "Chaque licence est unique et traçable (ton tableau de licences)",
    ])

    # ── 7. CONTACT ──────────────────────────────────────────────────────────
    h(doc, "7. Contact et support")
    para(doc,
        "Pour toute question sur les licences, la génération de clés ou le support client :")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("contact.amah.officiel@gmail.com")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = GOLD

    # Sauvegarde
    path = r"C:\Users\Smarte technologui\Desktop\Projets\amah-agent\developpeur\GUIDE_LICENCE_COMPLET.docx"
    doc.save(path)
    print(f"Document cree : {path}")
    os.startfile(path)


if __name__ == "__main__":
    main()
