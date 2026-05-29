"""
Génère le document de présentation du projet Amah Agent en Word.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

GOLD  = RGBColor(0xC8, 0xA9, 0x6E)
DARK  = RGBColor(0x1A, 0x1A, 0x17)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY  = RGBColor(0x55, 0x55, 0x55)


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p    = doc.add_paragraph()
    run  = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(20)
        run.font.color.rgb = GOLD
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        run.font.size  = Pt(14)
        run.font.color.rgb = GOLD
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    else:
        run.font.size  = Pt(12)
        run.font.color.rgb = DARK
        run.italic = True
        p.paragraph_format.space_before = Pt(8)
    return p


def add_para(doc, text, italic=False, bold=False, color=None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.italic     = italic
    run.bold       = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, indent=0):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3 * (indent + 1))
    p.paragraph_format.space_after = Pt(2)
    return p


def add_tool_table(doc, tools_data):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    for i, txt in enumerate(["Outil", "Paramètres", "Description"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_bg(hdr[i], '1A1A17')

    for name, params, desc in tools_data:
        row = table.add_row().cells
        row[0].text = name
        row[0].paragraphs[0].runs[0].bold = True
        row[0].paragraphs[0].runs[0].font.color.rgb = GOLD
        row[1].text = params
        row[1].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].paragraphs[0].runs[0].font.color.rgb = GREY
        row[2].text = desc

    doc.add_paragraph()


def main():
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    # ════════════════════════════════════════════════════════
    # PAGE DE TITRE
    # ════════════════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("THE AMAH — AGENT IA LOCAL")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = GOLD

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Documentation complète du projet")
    run2.font.size = Pt(14)
    run2.font.color.rgb = GREY

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = date_p.add_run(f"Version 1.0 — {datetime.now().strftime('%d/%m/%Y')}")
    run3.font.size = Pt(11)
    run3.font.color.rgb = GREY
    run3.italic = True

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════
    # 1. PRÉSENTATION DU PROJET
    # ════════════════════════════════════════════════════════
    add_heading(doc, "1. Présentation du projet", 1)

    add_para(doc,
        "Amah Agent est un assistant IA installé directement sur un PC Windows. "
        "Contrairement aux chatbots en ligne, il fonctionne localement et peut interagir "
        "avec le système de fichiers, envoyer des emails, naviguer sur internet et "
        "mémoriser des informations entre les sessions.")

    add_para(doc,
        "L'intelligence artificielle est fournie par Groq (modèle Llama 3.3-70B), "
        "un service gratuit offrant jusqu'à 14 400 requêtes par jour. "
        "L'agent dispose de 27 outils réels lui permettant d'agir concrètement "
        "sur le PC de l'utilisateur.")

    add_heading(doc, "Caractéristiques principales", 2)
    for item in [
        "27 outils réels (fichiers, documents, email, web, mémoire, navigateur)",
        "Interface graphique moderne — thème or/sombre, fenêtre Windows native",
        "Mémoire persistante SQLite — se souvient entre les sessions",
        "Email Gmail intégré (SMTP/IMAP) — lire, envoyer, chercher",
        "Navigation web automatique avec Chrome visible (Playwright)",
        "Packaging .exe — distributable sans installation de Python",
        "Email officiel : contact.amah.officiel@gmail.com",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════
    # 2. ARCHITECTURE TECHNIQUE
    # ════════════════════════════════════════════════════════
    add_heading(doc, "2. Architecture technique", 1)

    add_heading(doc, "Stack technologique", 2)
    for item in [
        "Langage : Python 3.13",
        "IA : Groq API — modèle llama-3.3-70b-versatile (gratuit, 30 req/min)",
        "Interface : tkinter (natif Python) avec thème or/sombre personnalisé",
        "Base de données : SQLite (natif Python) — fichier amah_memory.db",
        "Email : smtplib + imaplib (natif Python) — pas de dépendance Google Cloud",
        "Navigateur : Playwright + Chromium (Chrome headful, l'utilisateur voit l'écran)",
        "Packaging : PyInstaller --onefile → Amah Agent.exe (80 Mo)",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Structure des fichiers", 2)
    add_para(doc, "Le projet est organisé comme suit :")
    for item in [
        "agent.py — Interface terminal (Rich, couleurs or/sombre)",
        "gui.py — Interface graphique (écran de config + chat + mémoire)",
        "config.py — Modèle IA, system prompt, 27 définitions d'outils",
        "tools/files.py — 7 outils fichiers/dossiers",
        "tools/documents.py — 4 outils création/lecture documents",
        "tools/search.py — 2 outils recherche web DuckDuckGo",
        "tools/system.py — 3 outils système Windows",
        "tools/memory.py — 3 outils mémoire SQLite",
        "tools/email_tool.py — 3 outils Gmail SMTP/IMAP",
        "tools/browser.py — 5 outils navigation Playwright",
        "amah_memory.db — Base SQLite (créée automatiquement)",
        "amah.ico — Icône de l'application",
        "build.bat — Compilation en .exe en une commande",
        "dist/ — Dossier livrable client (exe + guide + scripts)",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════
    # 3. LES 27 OUTILS
    # ════════════════════════════════════════════════════════
    add_heading(doc, "3. Les 27 outils", 1)

    add_heading(doc, "Fichiers et dossiers (7 outils)", 2)
    add_tool_table(doc, [
        ("list_files",      "path",                    "Liste les fichiers et dossiers d'un répertoire"),
        ("organize_folder", "path",                    "Classe automatiquement les fichiers par type"),
        ("find_files",      "path, pattern",           "Cherche des fichiers par nom ou extension"),
        ("move_file",       "source, destination",     "Déplace ou renomme un fichier"),
        ("create_folder",   "path",                    "Crée un nouveau dossier"),
        ("read_file",       "path",                    "Lit le contenu d'un fichier texte"),
        ("get_folder_info", "path",                    "Statistiques : taille, nombre de fichiers, types"),
    ])

    add_heading(doc, "Documents (4 outils)", 2)
    add_tool_table(doc, [
        ("create_word",    "filename, title, content", "Crée un document Word (.docx)"),
        ("create_txt",     "filename, content",        "Crée un fichier texte (.txt)"),
        ("create_pdf",     "filename, title, content", "Crée un document PDF"),
        ("read_document",  "path",                     "Lit un fichier Word, PDF ou TXT"),
    ])

    add_heading(doc, "Recherche web (2 outils)", 2)
    add_tool_table(doc, [
        ("web_search",  "query, num_results", "Recherche DuckDuckGo — retourne les résultats"),
        ("read_webpage","url",                "Extrait le texte d'une page web depuis son URL"),
    ])

    add_heading(doc, "Système Windows (3 outils)", 2)
    add_tool_table(doc, [
        ("get_system_info", "—",     "RAM, CPU, disque, version Windows"),
        ("open_file",       "path",  "Ouvre un fichier ou dossier avec le programme par défaut"),
        ("run_command",     "command","Exécute une commande PowerShell (commandes dangereuses bloquées)"),
    ])

    add_heading(doc, "Mémoire persistante (3 outils)", 2)
    add_para(doc,
        "La mémoire utilise une base SQLite locale (amah_memory.db). "
        "Elle fonctionne sur deux niveaux : l'historique automatique des conversations "
        "(rechargé à chaque session) et la mémoire explicite que l'utilisateur demande de retenir.")
    add_tool_table(doc, [
        ("save_memory",   "content, category", "Mémorise une information (catégories : préférence, tâche, info, projet)"),
        ("get_memories",  "category",          "Rappelle les informations mémorisées"),
        ("delete_memory", "memory_id",         "Supprime une mémoire par son identifiant"),
    ])

    add_heading(doc, "Email Gmail (3 outils)", 2)
    add_para(doc,
        "L'email fonctionne via SMTP/IMAP avec un mot de passe d'application Google. "
        "Aucun Google Cloud ni OAuth2 requis. L'adresse officielle est "
        "contact.amah.officiel@gmail.com.")
    add_tool_table(doc, [
        ("read_emails",   "n",                    "Lit les N derniers emails de la boîte Gmail"),
        ("send_email",    "to, subject, body",    "Envoie un email depuis le compte Gmail configuré"),
        ("search_emails", "query",                "Cherche des emails (syntaxe Gmail : from:, subject:, is:unread...)"),
    ])

    add_heading(doc, "Navigation web Playwright (5 outils)", 2)
    add_para(doc,
        "Chrome s'ouvre de façon visible (headless=False) — l'utilisateur voit "
        "ce que fait l'agent en temps réel. Le navigateur reste ouvert entre les actions.")
    add_tool_table(doc, [
        ("open_browser",    "url",            "Ouvre une page web dans Chrome"),
        ("click_element",   "selector",       "Clique sur un élément (CSS selector ou texte)"),
        ("fill_form",       "selector, value","Remplit un champ de formulaire"),
        ("take_screenshot", "path",           "Capture une image de l'écran du navigateur"),
        ("get_page_text",   "—",              "Lit et retourne le contenu texte de la page ouverte"),
    ])

    # ════════════════════════════════════════════════════════
    # 4. MÉMOIRE ET CONTINUITÉ
    # ════════════════════════════════════════════════════════
    add_heading(doc, "4. Système de mémoire", 1)

    add_para(doc,
        "Amah Agent dispose d'une mémoire persistante SQLite qui survit entre les sessions. "
        "Elle fonctionne sur deux niveaux complémentaires :")

    add_heading(doc, "Niveau 1 — Historique automatique", 2)
    add_para(doc,
        "Chaque échange (message utilisateur + réponse d'Amah) est automatiquement "
        "sauvegardé dans la table 'conversations'. Au prochain démarrage, les 20 derniers "
        "messages sont rechargés dans le contexte — Amah reprend la conversation "
        "là où elle s'est arrêtée, sans que l'utilisateur ait à répéter.")

    add_heading(doc, "Niveau 2 — Mémoire long terme explicite", 2)
    add_para(doc,
        "L'utilisateur peut demander à Amah de mémoriser des informations importantes : "
        "préférences, habitudes, informations de contact, tâches récurrentes. "
        "Ces informations sont stockées dans la table 'memories' avec une catégorie "
        "et sont accessibles à tout moment via get_memories.")

    # ════════════════════════════════════════════════════════
    # 5. INTERFACE GRAPHIQUE
    # ════════════════════════════════════════════════════════
    add_heading(doc, "5. Interface graphique", 1)

    add_para(doc,
        "L'interface graphique (gui.py) est construite avec tkinter, "
        "la bibliothèque graphique native de Python. Elle utilise un thème "
        "personnalisé or/sombre cohérent avec l'identité visuelle d'Amah.")

    add_heading(doc, "Fonctionnalités de l'interface", 2)
    for item in [
        "Écran de configuration automatique au premier lancement (crée le .env)",
        "Zone de chat avec historique des échanges coloré (or pour Amah, blanc pour l'utilisateur)",
        "Indicateur 'réfléchit...' pendant les appels à l'IA",
        "Appels API en thread séparé — l'interface ne se bloque jamais",
        "Commandes spéciales : 'outils', 'reset'",
        "Chargement automatique des 20 derniers messages au démarrage",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Écran de configuration (premier lancement)", 2)
    add_para(doc,
        "La première fois que le client lance Amah Agent.exe, une fenêtre de configuration "
        "s'affiche automatiquement. Il doit renseigner sa clé API Groq (obligatoire) "
        "et optionnellement ses identifiants Gmail. Un fichier .env est créé automatiquement "
        "— le client ne touche jamais au code ou aux fichiers de configuration manuellement.")

    # ════════════════════════════════════════════════════════
    # 6. DISTRIBUTION ET PACKAGING
    # ════════════════════════════════════════════════════════
    add_heading(doc, "6. Distribution et packaging", 1)

    add_para(doc,
        "Amah Agent est distribué sous forme d'un fichier exécutable Windows (.exe) "
        "généré par PyInstaller. Le client n'a pas besoin d'installer Python "
        "ou aucune autre dépendance.")

    add_heading(doc, "Contenu du livrable client", 2)
    for item in [
        "Amah Agent.exe (80 Mo) — l'application complète, standalone",
        "installer_navigateur.bat — installe Chrome pour la navigation web (une seule fois)",
        "GUIDE_INSTALLATION.md — guide complet d'installation étape par étape",
        ".env — fichier de configuration avec la clé Groq et les identifiants Gmail",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Processus de build", 2)
    add_para(doc,
        "Pour générer un nouveau .exe après modification du code, "
        "il suffit de double-cliquer sur build.bat. Le script : "
        "nettoie les anciens fichiers, recompile avec PyInstaller, "
        "et copie automatiquement les fichiers de distribution dans le dossier dist/.")

    add_heading(doc, "Sécurité", 2)
    for item in [
        "La clé API Groq est stockée dans .env — jamais dans le code ni dans le .exe",
        "Le mot de passe Gmail est stocké dans .env — jamais dans le code",
        "Les commandes PowerShell dangereuses sont bloquées (rm -rf, format, shutdown...)",
        "Le .env.example ne contient jamais de vraies credentials",
        "Aucune donnée n'est envoyée vers des serveurs tiers sauf Groq pour l'IA",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════
    # 7. DÉPENDANCES
    # ════════════════════════════════════════════════════════
    add_heading(doc, "7. Dépendances Python", 1)

    add_tool_table(doc, [
        ("groq",                  "pip install groq",           "Client API Groq — connexion au modèle IA"),
        ("python-dotenv",         "pip install python-dotenv",  "Chargement des variables d'environnement (.env)"),
        ("rich",                  "pip install rich",           "Interface terminal colorée (agent.py)"),
        ("ddgs",                  "pip install ddgs",           "Recherche web DuckDuckGo"),
        ("python-docx",           "pip install python-docx",    "Création de documents Word"),
        ("fpdf2",                 "pip install fpdf2",          "Création de documents PDF"),
        ("pypdf",                 "pip install pypdf",          "Lecture de documents PDF"),
        ("psutil",                "pip install psutil",         "Informations système (RAM, CPU, disque)"),
        ("playwright",            "pip install playwright",     "Navigation web automatique (+ playwright install chromium)"),
        ("pyinstaller",           "pip install pyinstaller",    "Packaging en .exe Windows"),
        ("pillow",                "pip install pillow",         "Génération de l'icône .ico"),
    ])

    add_para(doc,
        "Bibliothèques natives Python utilisées (sans installation) : "
        "tkinter, sqlite3, smtplib, imaplib, email, pathlib, threading, json, os, sys.",
        italic=True, color=GREY)

    # ════════════════════════════════════════════════════════
    # 8. PROCHAINES ÉVOLUTIONS
    # ════════════════════════════════════════════════════════
    add_heading(doc, "8. Prochaines évolutions", 1)

    for item in [
        "Outil agenda Google Calendar — lire et créer des événements",
        "Synthèse vocale — Amah répond à voix haute",
        "Mise à jour automatique du .exe",
        "Mode multi-utilisateurs avec profils séparés",
        "Interface web locale (alternative à tkinter)",
        "Avatar animé dans l'interface graphique",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════
    # 9. CONTACT ET SUPPORT
    # ════════════════════════════════════════════════════════
    add_heading(doc, "9. Contact et support", 1)

    add_para(doc, "Email officiel Amah : contact.amah.officiel@gmail.com")
    add_para(doc, "Pour toute question technique ou demande de personnalisation, "
             "contactez l'équipe via l'adresse email officielle du projet.")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer.add_run("— THE AMAH — AGENT IA LOCAL — Documentation v1.0 —")
    run_f.italic = True
    run_f.font.color.rgb = GOLD
    run_f.font.size = Pt(10)

    # Sauvegarde
    path = r"C:\Users\Smarte technologui\Desktop\Projets\amah-agent\Amah_Agent_Documentation.docx"
    doc.save(path)
    print(f"Document cree : {path}")
    os.startfile(path)


if __name__ == "__main__":
    main()
